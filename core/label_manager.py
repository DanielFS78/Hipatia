"""
========================================================================
GESTOR DE ETIQUETAS - SISTEMA DE IMPRESIÓN CON PLANTILLAS WORD
========================================================================
Gestiona la creación e impresión de etiquetas usando plantillas Word.

Características:
- Lee plantillas .docx con placeholders
- Reemplaza {{producto}}, {{descripcion}}, {{codigo}}
- Inserta códigos QR en posiciones {{qr}}
- Soporta múltiples formatos de etiquetas (A5, A4)
- Genera documentos listos para imprimir

Formatos soportados:
- APLI 1857 (A5): 8x12mm - Para QR pequeños
- A4 Estándar: 105x42mm - 14 etiquetas por hoja

Autor: Sistema de Trazabilidad
Fecha: 2025
========================================================================
"""

import logging
import os
from pathlib import Path
from typing import Optional, Dict, List, Tuple
from datetime import datetime
import tempfile

try:
    from docx import Document
    from docx.shared import Inches, Mm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise ImportError(
        "python-docx no está instalado. Instala con: pip install python-docx"
    )

from PIL import Image
from io import BytesIO


class LabelManager:
    """
    Gestor de etiquetas basado en plantillas Word.

    Permite trabajar con diferentes formatos de etiquetas APLI
    usando plantillas .docx personalizables.

    Attributes:
        logger: Logger para registro de operaciones
        templates_dir: Directorio raíz de plantillas
        qr_generator: Instancia del generador de QR (opcional)
    """

    # Configuraciones de formatos conocidos
    LABEL_FORMATS = {
        'APLI_1857_A5': {
            'nombre': 'APLI 1857 (A5)',
            'tamaño_etiqueta': (8, 12),  # mm
            'formato_hoja': 'A5',
            'descripcion': 'Etiquetas pequeñas 8x12mm para QR mínimo'
        },
        'A4_14_ETIQUETAS': {
            'nombre': 'A4 - 14 etiquetas',
            'tamaño_etiqueta': (105, 42),  # mm
            'formato_hoja': 'A4',
            'etiquetas_por_hoja': 14,
            'descripcion': 'Etiquetas estándar 105x42mm'
        }
    }

    def __init__(
            self,
            templates_dir: str = "templates",
            qr_generator=None
    ):
        """
        Inicializa el gestor de etiquetas.

        Args:
            templates_dir: Directorio raíz de las plantillas
            qr_generator: Instancia de QrGenerator (opcional)
        """
        self.logger = logging.getLogger("EvolucionTiemposApp.LabelManager")
        self.templates_dir = Path(templates_dir)
        self.qr_generator = qr_generator

        # Crear estructura de directorios si no existe
        self._ensure_template_structure()

        self.logger.info(f"LabelManager inicializado. Templates: {self.templates_dir}")

    def _ensure_template_structure(self):
        """Crea la estructura de directorios de plantillas si no existe."""
        try:
            directories = [
                self.templates_dir / "etiquetas" / "A5",
                self.templates_dir / "etiquetas" / "A4",
                self.templates_dir / "documentos"
            ]

            for directory in directories:
                directory.mkdir(parents=True, exist_ok=True)
                self.logger.debug(f"Directorio verificado: {directory}")

        except Exception as e:
            self.logger.error(f"Error creando estructura de directorios: {e}")

    def get_template_path(
            self,
            formato: str,
            nombre_plantilla: str
    ) -> Optional[Path]:
        """
        Obtiene la ruta completa de una plantilla.

        Args:
            formato: 'A5', 'A4', o 'documentos'
            nombre_plantilla: Nombre del archivo de plantilla

        Returns:
            Path completo a la plantilla o None si no existe
        """
        if formato in ['A5', 'A4']:
            ruta = self.templates_dir / "etiquetas" / formato / nombre_plantilla
        else:
            ruta = self.templates_dir / formato / nombre_plantilla

        if ruta.exists():
            return ruta
        else:
            self.logger.warning(f"Plantilla no encontrada: {ruta}")
            return None

    def list_templates(self, formato: str = None) -> List[Dict]:
        """
        Lista todas las plantillas disponibles.

        Args:
            formato: Filtrar por formato (A5, A4, documentos) o None para todas

        Returns:
            Lista de diccionarios con info de plantillas
        """
        templates = []

        if formato:
            formatos = [formato]
        else:
            formatos = ['A5', 'A4', 'documentos']

        for fmt in formatos:
            if fmt in ['A5', 'A4']:
                path = self.templates_dir / "etiquetas" / fmt
            else:
                path = self.templates_dir / fmt

            if path.exists():
                for archivo in path.glob("*.docx"):
                    if not archivo.name.startswith('~'):  # Ignorar archivos temporales
                        templates.append({
                            'nombre': archivo.name,
                            'formato': fmt,
                            'ruta': str(archivo),
                            'tamaño': archivo.stat().st_size,
                            'modificado': datetime.fromtimestamp(archivo.stat().st_mtime)
                        })

        return templates

    def count_qr_placeholders(self, plantilla: str, formato: str) -> int:
        """
        Cuenta el número total de placeholders '{{qr}}' en una plantilla.

        Args:
            plantilla: Nombre del archivo de plantilla (ej: 'qr.docx')
            formato: Formato (ej: 'A5', 'A4')

        Returns:
            int: Número total de '{{qr}}' encontrados, o 0 si hay error.
        """
        self.logger.info(f"Contando placeholders '{{qr}}' en {formato}/{plantilla}...")

        try:
            # 1. Obtener la ruta de la plantilla
            template_path = self.get_template_path(formato, plantilla)
            if not template_path:
                self.logger.error(f"No se pudo encontrar la plantilla para contar: {plantilla}")
                return 0

            # 2. Cargar el documento
            doc = Document(str(template_path))
            count = 0

            # 3. Contar en párrafos
            for paragraph in doc.paragraphs:
                count += paragraph.text.count('{{qr}}')

            # 4. Contar en tablas (celda por celda)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            count += paragraph.text.count('{{qr}}')

            self.logger.info(f"Se encontraron {count} placeholders '{{qr}}' en la plantilla.")
            return count

        except Exception as e:
            self.logger.error(f"Error al contar placeholders '{{qr}}': {e}", exc_info=True)
            return 0

    def replace_placeholders(
            self,
            doc: Document,
            datos: Dict[str, str],
            insert_qr: bool = True
    ) -> Document:
        """
        Reemplaza placeholders en un documento Word.

        Placeholders soportados:
        - {{producto}} - Nombre del producto
        - {{descripcion}} - Descripción
        - {{codigo}} - Código del producto
        - {{qr}} - Posición para código QR
        - {{fecha}} - Fecha actual
        - Cualquier otro campo en el diccionario datos

        Args:
            doc: Documento Word (python-docx)
            datos: Diccionario con los datos a reemplazar
            insert_qr: Si True, inserta imagen QR en posiciones {{qr}}

        Returns:
            Documento modificado
        """
        try:
            # Añadir fecha actual si no está en datos
            if 'fecha' not in datos:
                datos['fecha'] = datetime.now().strftime('%d/%m/%Y')

            # Reemplazar en párrafos
            for paragraph in doc.paragraphs:
                self._replace_in_runs(paragraph.runs, datos, insert_qr)

            # Reemplazar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_runs(paragraph.runs, datos, insert_qr)

            self.logger.debug("Placeholders reemplazados correctamente")
            return doc

        except Exception as e:
            self.logger.error(f"Error reemplazando placeholders: {e}")
            return doc

    def _replace_in_runs(
            self,
            runs,
            datos: Dict[str, str],
            insert_qr: bool
    ):
        """
        Reemplaza placeholders en runs de texto.
        VERSIÓN CORREGIDA: Maneja múltiples '{{qr}}' en un mismo párrafo.
        """
        if not runs:
            return

        # 1. Concatenar todo el texto del párrafo
        texto_completo = ''.join(run.text for run in runs)

        # 2. Reemplazar primero los placeholders de TEXTO
        for key, value in datos.items():
            if key == 'qr':  # Saltamos 'qr' por ahora
                continue
            placeholder = f"{{{{{key}}}}}"
            if placeholder in texto_completo:
                texto_completo = texto_completo.replace(placeholder, str(value))

        # 3. Manejar el reemplazo de IMAGEN (QR)
        if insert_qr and '{{qr}}' in texto_completo:
            qr_data = datos.get('codigo', 'QR_ERROR')

            # Dividir el texto por el placeholder '{{qr}}'
            # Ej: "{{qr}},{{qr}}" se vuelve -> ["", ",", ""]
            parts = texto_completo.split('{{qr}}')

            # Limpiar todos los runs existentes y empezar desde el primero
            # Escribimos la primera parte (texto antes del primer QR)
            runs[0].text = parts[0]
            for run in runs[1:]:
                run.text = ''  # Vaciar todos los demás runs

            # 4. Reconstruir el párrafo alternando imagen y texto
            # Empezamos desde la segunda parte (parts[1])
            for part in parts[1:]:
                # a. Insertar la imagen QR
                self._insert_qr_image_to_run(runs[0], qr_data)

                # b. Añadir el texto que venía DESPUÉS de ese QR
                runs[0].add_text(part)

        else:
            # Si no había '{{qr}}', simplemente escribimos el texto (ya reemplazado)
            runs[0].text = texto_completo
            for run in runs[1:]:
                run.text = ''

    def _insert_qr_image_to_run(
            self,
            run, # Recibe el 'run' donde debe añadir la imagen
            qr_data: str
    ):
        """
        Inserta una imagen QR en un 'run' específico.
        No borra texto, solo AÑADE la imagen al final del 'run'.
        """
        try:
            if not self.qr_generator:
                self.logger.warning("No hay generador de QR configurado")
                return

            # Generar QR (300x300 da buena resolución para imprimir pequeño)
            qr_img = self.qr_generator.generate_qr_code(qr_data, size=(300, 300))
            if not qr_img:
                self.logger.error(f"No se pudo generar la imagen QR para: {qr_data}")
                return

            # Convertir PIL Image a bytes
            img_buffer = BytesIO()
            qr_img.save(img_buffer, format='PNG')
            img_buffer.seek(0)

            # Insertar imagen (el tamaño físico final sigue siendo 8mm)
            # IMPORTANTE: Esto AÑADE la imagen al final del contenido
            # actual del 'run'.
            run.add_picture(img_buffer, width=Mm(11))

            self.logger.debug(f"Imagen QR insertada para: {qr_data}")

        except Exception as e:
            self.logger.error(f"Error insertando imagen QR: {e}", exc_info=True)

    def generate_labels(
            self,
            plantilla: str,
            formato: str,
            datos_lista: List[Dict[str, str]],
            output_path: Optional[str] = None
    ) -> Optional[str]:
        """
        Genera un documento con múltiples etiquetas.
        MODIFICADO: Itera sobre 'datos_lista' y rellena los placeholders
        en secuencia (idealmente en las celdas de una tabla).

        Args:
            plantilla: Nombre de la plantilla a usar
            formato: Formato de etiqueta (A5, A4, documentos)
            datos_lista: Lista de diccionarios con datos para CADA etiqueta
            output_path: Ruta de salida (opcional, genera temporal si no se especifica)

        Returns:
            Ruta del archivo generado o None si hay error
        """
        try:
            # 1. Obtener ruta de plantilla
            template_path = self.get_template_path(formato, plantilla)
            if not template_path:
                self.logger.error(f"Plantilla no encontrada: {plantilla}")
                return None

            # 2. Cargar plantilla
            doc = Document(str(template_path))

            # --- INICIO DE LA NUEVA LÓGICA ---

            # 3. Crear un iterador de los datos
            #    Esto nos permite hacer data_iterator.next()
            data_iterator = iter(datos_lista)

            # 4. Asumir que las etiquetas están en tablas (lo más común)
            #    Recorremos cada celda y la tratamos como una etiqueta.

            stop_filling = False
            for table in doc.tables:
                if stop_filling:
                    break
                for row in table.rows:
                    if stop_filling:
                        break
                    for cell in row.cells:
                        try:
                            # 5. Tomar el SIGUIENTE item de la lista de datos
                            datos_para_esta_celda = next(data_iterator)

                            # 6. Reemplazar placeholders SOLO en esta celda
                            for paragraph in cell.paragraphs:
                                self._replace_in_runs(
                                    paragraph.runs,
                                    datos_para_esta_celda,
                                    insert_qr=True
                                )

                        except StopIteration:
                            # Se nos acabaron los datos (ej: 5 QRs en una hoja de 8)
                            # Dejamos de rellenar.
                            stop_filling = True
                            self.logger.info("Se han rellenado todos los datos. Dejando el resto de la plantilla.")
                            break
                        except Exception as e:
                            self.logger.warning(f"Error rellenando una celda: {e}")
                            continue  # Intentar con la siguiente celda

            # --- FIN DE LA NUEVA LÓGICA ---

            # 7. Guardar documento
            if not output_path:
                # Crear archivo temporal
                temp_dir = tempfile.gettempdir()
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_path = os.path.join(
                    temp_dir,
                    f"etiquetas_{formato}_{timestamp}.docx"
                )

            doc.save(output_path)
            self.logger.info(f"Etiquetas generadas: {output_path} (con {len(datos_lista)} etiquetas únicas)")
            return output_path

        except Exception as e:
            self.logger.error(f"Error generando etiquetas: {e}", exc_info=True)
            return None

    def _replicate_template_content(
            self,
            doc: Document,
            count: int
    ) -> Document:
        """
        Replica el contenido de una plantilla múltiples veces.

        Args:
            doc: Documento original
            count: Número de veces a replicar

        Returns:
            Documento con contenido replicado
        """
        # Nota: Esto es una simplificación
        # Para producción, necesitarías lógica más compleja
        # dependiendo de cómo esté estructurada tu plantilla

        # Por ahora, simplemente añade saltos de página
        for i in range(count - 1):
            doc.add_page_break()

        return doc

    def _is_printer_available(self) -> bool:
        """
        Comprueba si hay una impresora predeterminada configurada.

        Returns:
            True si hay impresora disponible, False si no.
        """
        import subprocess
        import platform

        system = platform.system()

        try:
            if system == 'Darwin':  # macOS
                result = subprocess.run(
                    ['lpstat', '-d'],
                    capture_output=True,
                    text=True,
                    timeout=5
                )
                # Si hay impresora, la salida contiene "system default destination:"
                # Si NO hay impresora, contiene "no system default destination"
                if 'no system default destination' in result.stdout.lower():
                    self.logger.info("No hay impresora predeterminada configurada.")
                    return False
                return True

            elif system == 'Linux':
                result = subprocess.run(
                    ['lpstat', '-d'],
                    capture_output=True,
                    text=True,
                    timeout=5
                )
                if 'no system default destination' in result.stdout.lower() or result.returncode != 0:
                    self.logger.info("No hay impresora predeterminada configurada.")
                    return False
                return True

            elif system == 'Windows':
                # En Windows, asumimos que hay impresora si no falla el sistema
                # (la detección es más compleja)
                return True

            else:
                return True  # Asumir disponible para otros sistemas

        except Exception as e:
            self.logger.warning(f"Error comprobando impresora: {e}")
            return False

    def _save_to_documents(self, doc_path: str) -> Optional[str]:
        """
        Guarda el documento en la carpeta de Documentos del usuario.

        Args:
            doc_path: Ruta del archivo temporal a copiar.

        Returns:
            Ruta del archivo guardado o None si hay error.
        """
        import shutil
        from pathlib import Path

        try:
            # Crear carpeta ~/Documents/Etiquetas/
            home = Path.home()
            etiquetas_dir = home / "Documents" / "Etiquetas"
            etiquetas_dir.mkdir(parents=True, exist_ok=True)

            # Copiar archivo
            filename = Path(doc_path).name
            dest_path = etiquetas_dir / filename

            shutil.copy2(doc_path, dest_path)
            self.logger.info(f"Documento guardado en: {dest_path}")

            return str(dest_path)

        except Exception as e:
            self.logger.error(f"Error guardando documento: {e}")
            return None

    def _open_file_location(self, file_path: str):
        """
        Abre la ubicación del archivo en el explorador de archivos.

        Args:
            file_path: Ruta del archivo a mostrar.
        """
        import subprocess
        import platform

        system = platform.system()

        try:
            if system == 'Darwin':  # macOS
                subprocess.run(['open', '-R', file_path])
            elif system == 'Windows':
                subprocess.run(['explorer', '/select,', file_path])
            elif system == 'Linux':
                subprocess.run(['xdg-open', str(Path(file_path).parent)])

        except Exception as e:
            self.logger.warning(f"No se pudo abrir la ubicación del archivo: {e}")

    def print_document(
            self,
            doc_path: str,
            printer_name: Optional[str] = None
    ) -> Tuple[bool, Optional[str]]:
        """
        Envía un documento a imprimir o lo guarda si no hay impresora.

        Args:
            doc_path: Ruta del documento a imprimir
            printer_name: Nombre de la impresora (None = predeterminada)

        Returns:
            Tupla (éxito, ruta_guardada):
            - (True, None) si se imprimió correctamente
            - (True, ruta) si se guardó porque no hay impresora
            - (False, None) si hubo error
        """
        try:
            import subprocess
            import platform

            # 1. Comprobar si hay impresora disponible
            if not self._is_printer_available():
                self.logger.info("No hay impresora. Guardando documento en Documentos...")
                saved_path = self._save_to_documents(doc_path)
                if saved_path:
                    self._open_file_location(saved_path)
                    return (True, saved_path)
                else:
                    return (False, None)

            # 2. Si hay impresora, imprimir normalmente
            system = platform.system()

            if system == 'Windows':
                if printer_name:
                    os.startfile(doc_path, 'print')
                else:
                    os.startfile(doc_path, 'print')

                self.logger.info(f"Documento enviado a imprimir: {doc_path}")
                return (True, None)

            elif system == 'Linux':
                cmd = ['lp']
                if printer_name:
                    cmd.extend(['-d', printer_name])
                cmd.append(doc_path)

                subprocess.run(cmd, check=True)
                self.logger.info(f"Documento enviado a imprimir: {doc_path}")
                return (True, None)

            elif system == 'Darwin':  # macOS
                cmd = ['lpr']
                if printer_name:
                    cmd.extend(['-P', printer_name])
                cmd.append(doc_path)

                subprocess.run(cmd, check=True)
                self.logger.info(f"Documento enviado a imprimir: {doc_path}")
                return (True, None)

            else:
                self.logger.error(f"Sistema operativo no soportado: {system}")
                return (False, None)

        except Exception as e:
            self.logger.error(f"Error al imprimir: {e}")
            return (False, None)

    def create_sample_template(
            self,
            formato: str,
            nombre: str = "plantilla_ejemplo.docx"
    ) -> Optional[str]:
        """
        Crea una plantilla de ejemplo.

        Args:
            formato: A5, A4 o documentos
            nombre: Nombre del archivo de plantilla

        Returns:
            Ruta de la plantilla creada o None si hay error
        """
        try:
            doc = Document()

            # Configurar según formato
            if formato == 'A5':
                # Plantilla para APLI 1857 (etiquetas pequeñas)
                doc.add_heading('Etiqueta APLI 1857 (A5)', level=2)
                doc.add_paragraph('Producto: {{producto}}')
                doc.add_paragraph('Código: {{codigo}}')
                doc.add_paragraph('QR: {{qr}}')

            elif formato == 'A4':
                # Plantilla para A4 estándar (14 etiquetas)
                table = doc.add_table(rows=7, cols=2)
                table.style = 'Table Grid'

                for row in table.rows:
                    for cell in row.cells:
                        cell.text = (
                            "{{producto}}\n"
                            "{{descripcion}}\n"
                            "Código: {{codigo}}\n"
                            "QR: {{qr}}"
                        )

            else:
                # Plantilla de documento genérico
                doc.add_heading('Documento de Fabricación', level=1)
                doc.add_paragraph(f'Fecha: {{fecha}}')
                doc.add_paragraph(f'Producto: {{producto}}')
                doc.add_paragraph(f'Descripción: {{descripcion}}')
                doc.add_paragraph(f'Código: {{codigo}}')

            # Guardar plantilla
            if formato in ['A5', 'A4']:
                output_path = self.templates_dir / "etiquetas" / formato / nombre
            else:
                output_path = self.templates_dir / formato / nombre

            doc.save(str(output_path))
            self.logger.info(f"Plantilla de ejemplo creada: {output_path}")
            return str(output_path)

        except Exception as e:
            self.logger.error(f"Error creando plantilla de ejemplo: {e}")
            return None


# ============================================================================
# FUNCIONES DE UTILIDAD
# ============================================================================

def quick_print_labels(
        datos: Dict[str, str],
        formato: str = 'A4',
        plantilla: str = 'plantilla_ejemplo.docx'
) -> bool:
    """
    Función rápida para imprimir una etiqueta.

    Args:
        datos: Diccionario con los datos de la etiqueta
        formato: A5 o A4
        plantilla: Nombre de la plantilla

    Returns:
        True si se imprimió correctamente
    """
    manager = LabelManager()
    doc_path = manager.generate_labels(plantilla, formato, [datos])

    if doc_path:
        return manager.print_document(doc_path)

    return False


# ============================================================================
# EJEMPLO DE USO
# ============================================================================

if __name__ == "__main__":
    # Configurar logging
    logging.basicConfig(level=logging.INFO)

    print("=" * 70)
    print("GESTOR DE ETIQUETAS - Ejemplo de Uso")
    print("=" * 70)

    # Crear gestor
    manager = LabelManager()

    # 1. Listar plantillas disponibles
    print("\n1. Plantillas disponibles:")
    templates = manager.list_templates()
    if templates:
        for t in templates:
            print(f"   - {t['formato']}/{t['nombre']}")
    else:
        print("   No hay plantillas. Creando ejemplos...")

        # Crear plantillas de ejemplo
        manager.create_sample_template('A5', 'apli_1857_ejemplo.docx')
        manager.create_sample_template('A4', 'a4_14_etiquetas_ejemplo.docx')
        print("   ✓ Plantillas de ejemplo creadas")

    # 2. Datos de ejemplo
    print("\n2. Preparando datos de ejemplo...")
    datos_etiqueta = {
        'producto': 'Widget ABC',
        'descripcion': 'Widget de alta precisión',
        'codigo': 'WDG-001',
        'qr': 'FAB123-WDG001-UNIT1-20250131-A3F9'
    }

    # 3. Generar etiquetas
    print("\n3. Generando etiquetas...")
    doc_path = manager.generate_labels(
        plantilla='a4_14_etiquetas_ejemplo.docx',
        formato='A4',
        datos_lista=[datos_etiqueta],
        output_path='etiquetas_test.docx'
    )

    if doc_path:
        print(f"   ✓ Etiquetas generadas: {doc_path}")
        print("\n4. ¿Deseas abrir el archivo? (s/n): ", end='')

        try:
            respuesta = input().lower()
            if respuesta == 's':
                import os

                os.startfile(doc_path) if os.name == 'nt' else os.system(f'open "{doc_path}"')
        except:
            pass

    print("\n" + "=" * 70)
    print("Ejemplo completado")
    print("=" * 70)