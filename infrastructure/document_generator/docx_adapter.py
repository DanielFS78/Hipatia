"""
Nombre del Módulo: infrastructure.document_generator.docx_adapter
Descripcion: Adaptador de infraestructura para la gestión de documentos usando la librería python-docx.
"""

import logging
import types
from typing import Optional, Dict, List, Any
from pathlib import Path
from io import BytesIO

docx: Optional[types.ModuleType] = None
Mm: Any = None

try:
    import docx as _docx_module
    from docx.shared import Mm as _Mm
    docx = _docx_module
    Mm = _Mm
except ImportError:
    pass

class DocxGeneratorAdapter:
    """
    Implementación del contrato IDocumentGenerator utilizando python-docx.

    Encapsula todo el acoplamiento con la librería externa, permitiendo
    la generación de documentos a partir de plantillas Word.
    """

    def __init__(self) -> None:
        self.logger = logging.getLogger("EvolucionTiemposApp.Infrastructure.DocxAdapter")
        if docx is None:
            self.logger.warning("La librería python-docx no está instalada.")

    def count_qr_placeholders(self, template_path: Path) -> int:
        """
        Cuenta los placeholders {{qr}} en la plantilla.

        Busca en párrafos y tablas del documento la cadena exacta '{{qr}}'.

        Args:
            template_path: Ruta al archivo .docx de la plantilla.

        Returns:
            Número total de menciones detectadas.
        """
        if not docx: return 0
        try:
            doc = docx.Document(str(template_path))
            count = 0
            for p in doc.paragraphs:
                count += p.text.count('{{qr}}')
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            count += p.text.count('{{qr}}')
            return count
        except Exception as e:
            self.logger.error(f"Error contando QR en {template_path}: {e}")
            return 0

    def generate_labels(
        self,
        template_path: Path,
        datos_lista: List[Dict[str, str]],
        qr_generator: Any,
        output_path: str,
        qr_size_mm: int = 11
    ) -> Optional[str]:
        """
        Genera un nuevo documento rellenando los datos en la plantilla.

        Args:
            template_path: Ruta a la plantilla base.
            datos_lista: Lista de diccionarios con la información a insertar.
            qr_generator: Objeto encargado de la creación de códigos QR.
            output_path: Ruta de destino para el archivo resultante.
            qr_size_mm: Tamaño lateral del código QR en mm.

        Returns:
            Ruta del archivo generado o None en caso de error.
        """
        if not docx: return None
        try:
            doc = docx.Document(str(template_path))
            data_iterator = iter(datos_lista)
            stop = False

            for table in doc.tables:
                if stop: break
                for row in table.rows:
                    if stop: break
                    for cell in row.cells:
                        try:
                            datos = next(data_iterator)
                            for p in cell.paragraphs:
                                self._replace_in_runs(p.runs, datos, True, qr_generator, qr_size_mm)
                        except StopIteration:
                            stop = True
                            break
                        except Exception:
                            continue

            doc.save(output_path)
            return output_path
        except Exception as e:
            self.logger.error(f"Error generando documento: {e}")
            return None

    def create_sample(self, format_name: str, output_path: Path) -> Optional[str]:
        """
        Crea un documento de ejemplo para un formato específico.

        Args:
            format_name: Identificador del formato (A5, APLI_1861_A5, A4, etc.).
            output_path: Ruta donde guardar el ejemplo.

        Returns:
            Ruta del archivo guardado o None.
        """
        if not docx: return None
        try:
            doc = docx.Document()
            if format_name == 'A5' or format_name == 'APLI_1861_A5':
                # Configurar sección A5
                section = doc.sections[0]
                section.page_width = Mm(148)
                section.page_height = Mm(210)
                section.top_margin = Mm(15)
                section.bottom_margin = Mm(15)
                section.left_margin = Mm(8)
                section.right_margin = Mm(8)
                
                # Crear tabla de muestra 11x6
                table = doc.add_table(rows=6, cols=11)
                table.style = 'Table Grid'
                table.alignment = 1 # Center
                
                # Eliminar el párrafo vacío inicial
                if len(doc.paragraphs) > 0:
                    p = doc.paragraphs[0]
                    p._element.getparent().remove(p._element)
                
                for row in table.rows:
                    row.height = Mm(30)
                    for cell in row.cells:
                        cell.width = Mm(12)
                        p = cell.paragraphs[0]
                        p.text = "{{qr}}"
            elif format_name == 'A4':
                table = doc.add_table(rows=7, cols=2)
                table.style = 'Table Grid'
                for r in table.rows:
                    for c in r.cells:
                        c.text = "{{producto}}\\n{{descripcion}}\\nCódigo: {{codigo}}\\nQR: {{qr}}"
            else:
                doc.add_heading('Documento de Fabricación', level=1)
                for f in ['fecha', 'producto', 'descripcion', 'codigo']:
                    doc.add_paragraph(f'{{{{{f}}}}}: {{{{{f}}}}}')

            doc.save(str(output_path))
            return str(output_path)
        except Exception as e:
            self.logger.error(f"Error creando ejemplo: {e}")
            return None

    # Métodos privados movivos de generator.py
    def _insert_qr_image_to_run(self, run: Any, qr_data: str, qr_generator: Any, qr_size_mm: int = 11) -> None:
        try:
            if not qr_generator: return
            qr_img = qr_generator.generate_qr_code(qr_data, size=(300, 300))
            if not qr_img: return
            img_buffer = BytesIO()
            qr_img.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            run.add_picture(img_buffer, width=Mm(qr_size_mm))
        except Exception as e:
            self.logger.error(f"Error insertando imagen QR: {e}")

    def _replace_in_runs(self, runs: Any, datos: Dict[str, str], insert_qr: bool, qr_generator: Any, qr_size_mm: int = 11) -> None:
        if not runs: return
        texto_completo = ''.join(run.text for run in runs)
        for key, value in datos.items():
            if key == 'qr': continue
            texto_completo = texto_completo.replace(f"{{{{{key}}}}}", str(value))
        
        if insert_qr and '{{qr}}' in texto_completo:
            qr_data = datos.get('codigo', 'QR_ERROR')
            parts = texto_completo.split('{{qr}}')
            runs[0].text = parts[0]
            for run in runs[1:]: run.text = ''
            for part in parts[1:]:
                self._insert_qr_image_to_run(runs[0], qr_data, qr_generator, qr_size_mm)
                runs[0].add_text(part)
        else:
            runs[0].text = texto_completo
            for run in runs[1:]: run.text = ''
