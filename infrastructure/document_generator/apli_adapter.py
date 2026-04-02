"""
Nombre del Módulo: infrastructure.document_generator.apli_adapter
Descripcion: Adaptador especializado para la generación de etiquetas APLI 01861 (A5, 66 etiquetas).
"""

from pathlib import Path
from typing import List, Dict, Any, Optional
import logging

docx_module: Any | None = None

try:
    import docx as _docx_module
    docx_module = _docx_module
    from docx.shared import Mm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from io import BytesIO
except ImportError:
    pass

from core.label_manager.ports import IDocumentGenerator

class Apli1861LabelGenerator(IDocumentGenerator):
    """
    Generador especializado para pegatinas APLI 01861 (12x30mm).

    Construye el documento A5 programáticamente con una tabla de 11 columnas
    por 6 filas, asegurando un ajuste preciso para las 66 etiquetas.
    """

    def __init__(self):
        self.logger = logging.getLogger("EvolucionTiemposApp.Apli1861LabelGenerator")

    def count_qr_placeholders(self, template_path: Path) -> int:
        """
        Calcula el número de huecos disponibles en este formato.

        Para este generador especializado, el número de huecos es fijo (66).
        Ignora el template_path ya que se genera dinámicamente.

        Args:
            template_path: Ruta a la plantilla (ignorada en este adaptador).

        Returns:
            Número total de etiquetas por hoja (66).
        """
        return 66

    def generate_labels(
        self,
        template_path: Path,
        datos_lista: List[Dict[str, str]],
        qr_generator: Any,
        output_path: str,
        qr_size_mm: int = 11
    ) -> Optional[str]:
        """
        Genera el documento A5 con el layout 11x6 de APLI 01861.

        Args:
            template_path: Ruta a la plantilla (ignorada).
            datos_lista: Lista de diccionarios con datos de cada etiqueta.
            qr_generator: Objeto encargado de generar las imágenes QR.
            output_path: Ruta donde se guardará el archivo generado.
            qr_size_mm: Tamaño del código QR en milímetros (default 11).

        Returns:
            La ruta del archivo generado o None si hubo un error.
        """
        if docx_module is None:
            self.logger.error("python-docx no está instalado.")
            return None

        try:
            doc = docx_module.Document()
            self.logger.info("CREANDO DOCUMENTO A5 PARA APLI 1861 (Nuclear Fix)")
            
            # Configurar página A5 Portrait
            section = doc.sections[0]
            section.page_width = Mm(148)
            section.page_height = Mm(210)
            section.orientation = WD_ORIENT.PORTRAIT
            
            # Ajustar márgenes segun especificaciones para 11x6 (14.5mm vertical para segurida, 8mm lateral)
            section.top_margin = Mm(14.5)
            section.bottom_margin = Mm(14.5)
            section.left_margin = Mm(8)
            section.right_margin = Mm(8)
            section.header_distance = Mm(0)
            section.footer_distance = Mm(0)

            # Eliminar párrafos vacíos iniciales que crea python-docx (para evitar saltos de página)
            for p in doc.paragraphs:
                p._element.getparent().remove(p._element)

            # --- MARCADOR ANTIGRAVITY v3 ---
            self.logger.info("Forzando tabla 6 filas x 11 columnas (66 etiquetas)...")
            table = doc.add_table(rows=6, cols=11)
            table.autofit = False 
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Datos a iterar
            data_iterator = iter(datos_lista)
            
            for row in table.rows:
                # 6 filas x 30mm = 180mm. Con 15mm de margen superior/inferior = 210mm (A5)
                row.height = Mm(30) 
                for cell in row.cells:
                    # 11 columnas x 12mm = 132mm. Con 8mm de margen lateral = 148mm (A5)
                    cell.width = Mm(12)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # Eliminar márgenes internos de la celda para maximizar espacio
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcMar = OxmlElement('w:tcMar')
                    for mar in ['top', 'left', 'bottom', 'right']:
                        node = OxmlElement(f'w:{mar}')
                        node.set(qn('w:w'), '0')
                        node.set(qn('w:type'), 'dxa')
                        tcMar.append(node)
                    tcPr.append(tcMar)
                    
                    try:
                        datos = next(data_iterator)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Si hay un generador de QR, insertamos la imagen
                        if qr_generator and "qr" in datos:
                            qr_data = datos["qr"]
                            try:
                                qr_img = qr_generator.generate_qr_code(qr_data, size=(300, 300))
                                if qr_img:
                                    img_buffer = BytesIO()
                                    qr_img.save(img_buffer, format='PNG')
                                    img_buffer.seek(0)
                                    run = p.add_run()
                                    run.add_picture(img_buffer, width=Mm(qr_size_mm), height=Mm(qr_size_mm))
                            except Exception as qr_err:
                                self.logger.error(f"Error insertando QR individual: {qr_err}")
                        else:
                            # Si no hay generador, ponemos texto informativo si existe
                            p.add_run(datos.get("codigo", ""))
                            
                    except StopIteration:
                        # No hay más datos, dejar celda vacía
                        continue

            doc.save(str(output_path))
            return str(output_path)

        except Exception as e:
            self.logger.error(f"Error generando etiquetas APLI 1861: {e}")
            import traceback
            self.logger.debug(traceback.format_exc())
            return None

    def create_sample(self, format_name: str, output_path: Path) -> Optional[str]:
        """Genera una muestra con placeholders."""
        # Podemos reutilizar la lógica de generación con datos de ejemplo
        sample_data = [{"qr": "PLACEHOLDER", "codigo": f"ETIQUETA {i+1}"} for i in range(66)]
        return self.generate_labels(Path("dynamic"), sample_data, None, str(output_path))
